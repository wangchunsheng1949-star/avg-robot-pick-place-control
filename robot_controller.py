#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
机器人取放货任务控制器
======================
基于JSON配置文件的取放货自动化系统

适配说明：
- 支持EXTERNAL类型动作（TAKE/PUT）
- 支持MOVE类型动作
- 支持instructions数组直接作为任务序列
"""

import sys
import time
import gc
import json
import math
import os
from datetime import timedelta, datetime
from typing import Dict, Any, List, Tuple, Optional, Callable
from dataclasses import dataclass
from pathlib import Path

# 第三方库
try:
    from kubo import KuboOrder, KuboInfo
except ImportError:
    # 模拟KuboOrder用于测试
    class KuboOrder:
        def __init__(self, ip, port, timeout=120, no_recv=False):
            self.ip = ip
            self.port = port
            self.timeout = timeout

        def robot_move(self, robot_id, pos, lift, ext, zoneType=1, maxSpeed=1.0):
            """移动机器人"""
            print(f"  [MOCK] 移动到 ({pos['x']}, {pos['y']}), 举升: {lift}mm")
            return True, {"status": "ok"}

        def robot_external(self, robot_id, action, position, height,
                           locationType, binType, binModel, timeout=60):
            """
            执行外部动作（取放货）
            action: "TAKE" 或 "PUT"
            """
            print(f"  [MOCK] 外部动作: {action}, 高度: {height}mm")
            return True, {"status": "ok", "action": action}

        def get_info(self, robot_id):
            """获取机器人状态"""
            return {
                "batteryInfo": {"powerLevel": 50},
                "errorState": []
            }

        def robot_init(self, robot_id):
            """初始化机器人"""
            pass

        def robot_pause(self, robot_id, pause_type):
            """暂停机器人"""
            pass

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ============================================================================
# 系统配置
# ============================================================================

@dataclass(frozen=True)
class SystemConfig:
    """系统级配置 - 从环境变量读取敏感信息"""
    IP: str = os.getenv("ROBOT_IP", "127.0.0.1")  # 机器人IP
    PORT: str = os.getenv("ROBOT_PORT", "9092")  # 机器人端口
    ROBOT_ID: str = os.getenv("ROBOT_ID", "ROBOT_001")  # 机器人ID
    DEFAULT_WORD_PATH: str = os.getenv("REPORT_PATH", "./report.docx")  # 报告路径
    DEFAULT_JSON_PATH: str = os.getenv("CONFIG_PATH", "./task_config.json")  # 配置路径


# ============================================================================
# JSON配置加载器
# ============================================================================

class TaskConfigLoader:
    """任务配置加载器 - 适配取放货JSON格式"""

    def __init__(self, json_path: str):
        self.json_path = json_path
        self.raw_config = self._load_json()
        self._validate()

        # 缓存
        self._locations_cache = {}
        self._instructions_cache = []

    def _load_json(self) -> Dict[str, Any]:
        """加载JSON文件"""
        try:
            with open(self.json_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            raise FileNotFoundError(f"配置文件未找到: {self.json_path}")
        except json.JSONDecodeError as e:
            raise ValueError(f"JSON格式错误: {e}")

    def _validate(self):
        """验证配置完整性"""
        # 检查必要字段
        if 'locations' not in self.raw_config:
            raise ValueError("缺少locations定义")
        if 'instructions' not in self.raw_config:
            raise ValueError("缺少instructions定义")

        # 验证充电点
        if 'charging_point' not in self.raw_config['locations']:
            raise ValueError("缺少charging_point定义")

    def get_metadata(self) -> Dict[str, Any]:
        """获取元数据"""
        return {
            "name": self.raw_config.get("task_name", "未命名"),
            "version": self.raw_config.get("version", "1.0"),
            "description": self.raw_config.get("description", ""),
            "loop": self.raw_config.get("loop", True)
        }

    def get_locations(self) -> Dict[str, Dict[str, Any]]:
        """
        获取所有位置定义
        返回格式: {name: {x, y, theta, zoneType, ...}}
        """
        if not self._locations_cache:
            self._locations_cache = self.raw_config.get('locations', {})
        return self._locations_cache

    def get_location(self, name: str) -> Dict[str, Any]:
        """获取指定位置的完整配置"""
        locations = self.get_locations()
        if name not in locations:
            raise ValueError(f"未知位置: {name}")
        return locations[name]

    def get_location_coord(self, name: str) -> Tuple[float, float, float]:
        """获取位置坐标 (x, y, theta)"""
        loc = self.get_location(name)
        return (loc['x'], loc['y'], loc['theta'])

    def get_instructions(self) -> List[Dict[str, Any]]:
        """
        获取指令列表（主任务序列）
        这是取放货JSON的核心
        """
        if not self._instructions_cache:
            self._instructions_cache = self.raw_config.get('instructions', [])
        return self._instructions_cache

    def get_power_strategy(self) -> Dict[str, Any]:
        """获取电源策略配置"""
        return self.raw_config.get('power_strategy', {
            "low_threshold": 20,
            "target_level": 80,
            "emergency_threshold": 5
        })

    def get_lift_config(self) -> Dict[str, float]:
        """获取举升配置"""
        return self.raw_config.get('lift_config', {
            "min_height": 600.0,
            "max_height": 8000.0,
            "speed": 500.0
        })

    def get_control_logic(self) -> Dict[str, Any]:
        """获取控制逻辑配置"""
        # 优先使用power_strategy中的值
        power = self.get_power_strategy()
        control = self.raw_config.get('control_logic', {})

        return {
            "max_runtime_hours": control.get('max_runtime_hours', 8),
            "battery_low_threshold": power.get('low_threshold', 20),
            "battery_high_threshold": power.get('target_level', 80),
            "battery_emergency_threshold": power.get('emergency_threshold', 5),
            "word_log_interval_hours": control.get('word_log_interval_hours', 2),
            "cycle_pause_seconds": control.get('cycle_pause_seconds', 2.0),
            "max_retry": control.get('max_retry', 3),
            "retry_delay": control.get('retry_delay', 2.0),
            "command_timeout": control.get('command_timeout', 120),
            "external_action_timeout": control.get('external_action_timeout', 60),
            "check_interval": power.get('check_interval', 10),
            "report_interval": power.get('report_interval', 1800)
        }


# ============================================================================
# 机器人指令执行器
# ============================================================================

class RobotCommandExecutor:
    """
    机器人指令执行器
    支持：
    - MOVE: 纯移动
    - EXTERNAL TAKE: 取货
    - EXTERNAL PUT: 放货
    """

    def __init__(self, robot: KuboOrder, robot_id: str,
                 task_config: TaskConfigLoader):
        self.robot = robot
        self.robot_id = robot_id
        self.config = task_config

        # 配置缓存
        self.locations = task_config.get_locations()
        self.lift_config = task_config.get_lift_config()
        self.control = task_config.get_control_logic()

        # 动作处理器映射
        self.action_handlers = {
            'MOVE': self._execute_move,
            'EXTERNAL': self._execute_external,
        }

        # 统计
        self.stats = {
            'chassis_distance': 0.0,
            'chassis_time': 0.0,
            'lift_distance': 0.0,
            'lift_time': 0.0,
            'pick_count': 0,
            'place_count': 0,
            'cycles': 0,
            'charge_count': 0
        }

        # 当前位置（用于距离计算）
        self._current_pos = task_config.get_location_coord('working_point')
        self._current_battery = 100

    def execute_instruction(self, instruction: Dict[str, Any]) -> bool:
        """
        执行单条instruction

        Args:
            instruction: JSON中的指令对象
        """
        inst_type = instruction.get('type')
        inst_id = instruction.get('id', 'unknown')

        self._log("Execute", f"[ID:{inst_id}] {instruction.get('description', inst_type)}")

        if inst_type not in self.action_handlers:
            self._log("Execute", f"未知指令类型: {inst_type}", "ERROR")
            return False

        handler = self.action_handlers[inst_type]

        try:
            success = handler(instruction)

            # 更新循环计数
            if success:
                self.stats['cycles'] += 1

            return success

        except Exception as e:
            self._log("Execute", f"执行异常: {e}", "ERROR")
            return False

    def _execute_move(self, instruction: Dict[str, Any]) -> bool:
        """执行MOVE移动指令"""
        location_ref = instruction.get('location_ref')
        if not location_ref:
            self._log("Move", "缺少location_ref", "ERROR")
            return False

        try:
            target = self.config.get_location_coord(location_ref)
            target_config = self.config.get_location(location_ref)
        except ValueError as e:
            self._log("Move", str(e), "ERROR")
            return False

        distance = self._calc_distance(self._current_pos, target)
        zone_type = target_config.get('zoneType', 0)
        lift_height = self.lift_config['min_height']

        success = self._do_move(target, lift_height, zone_type, 1.0)

        if success:
            self.stats['chassis_distance'] += distance
            self.stats['chassis_time'] += (distance / 1.0) / 3600
            self._current_pos = target

        return success

    def _execute_external(self, instruction: Dict[str, Any]) -> bool:
        """执行取放货EXTERNAL工艺指令"""
        action = instruction.get('action')
        position = instruction.get('position')
        height = instruction.get('height')

        if not all([action, position, height]):
            self._log("External", "缺少必要参数", "ERROR")
            return False

        target = (position['x'], position['y'], position['theta'])
        distance = self._calc_distance(self._current_pos, target)
        self._log("External", f"步骤1: 移动到作业精准点位")

        location_ref = instruction.get('location_ref', '')
        try:
            loc_config = self.config.get_location(location_ref)
            zone_type = loc_config.get('zoneType', 0)
        except:
            zone_type = 0

        if not self._do_move(target, height, zone_type, 0.5):
            return False
        self._current_pos = target
        self.stats['chassis_distance'] += distance
        self.stats['chassis_time'] += (distance / 0.5) / 3600

        self._log("External", f"步骤2: 执行{action}仓储装卸动作")
        location_type = instruction.get('locationType', 0)
        bin_type = instruction.get('binType', 10)
        bin_model = instruction.get('binModel', 'STANDARD0')
        timeout = self.control.get('external_action_timeout', 60)

        success = self._do_external(
            action, position, height,
            location_type, bin_type, bin_model, timeout
        )

        if success:
            if action == 'TAKE':
                self.stats['pick_count'] += 1
            elif action == 'PUT':
                self.stats['place_count'] += 1

            lift_dist = abs(height - self.lift_config['min_height']) / 1000.0
            self.stats['lift_distance'] += lift_dist
            self.stats['lift_time'] += (lift_dist * 1000 / self.lift_config['speed']) / 3600

        return success

    def _do_move(self, target: Tuple[float, float, float],
                 lift_height: float, zone_type: int,
                 max_speed: float) -> bool:
        """底层底盘运动执行"""
        max_retry = self.control['max_retry']
        retry_delay = self.control['retry_delay']

        for attempt in range(max_retry):
            try:
                success, result = self.robot.robot_move(
                    self.robot_id,
                    {'x': target[0], 'y': target[1], 'theta': target[2]},
                    lift_height,
                    0.0,
                    zoneType=zone_type,
                    maxSpeed=max_speed
                )

                if success:
                    return True
                else:
                    if attempt < max_retry - 1:
                        time.sleep(retry_delay)

            except Exception as e:
                self._log("Move", f"通讯异常: {e}", "WARN")
                if attempt < max_retry - 1:
                    time.sleep(retry_delay)

        return False

    def _do_external(self, action: str, position: Dict, height: float,
                     location_type: int, bin_type: int, bin_model: str,
                     timeout: int) -> bool:
        """底层仓储装卸工艺执行"""
        max_retry = 2
        for attempt in range(max_retry):
            try:
                success, result = self.robot.robot_external(
                    self.robot_id,
                    action,
                    position,
                    height,
                    location_type,
                    bin_type,
                    bin_model,
                    timeout
                )

                if success:
                    return True
                else:
                    if attempt < max_retry - 1:
                        time.sleep(1)

            except Exception as e:
                self._log("External", f"动作异常: {e}", "WARN")
                if attempt < max_retry - 1:
                    time.sleep(1)

        return False

    def _calc_distance(self, pos1: Tuple[float, ...], pos2: Tuple[float, ...]) -> float:
        """点位距离解算"""
        return math.sqrt((pos1[0] - pos2[0]) ** 2 + (pos1[1] - pos2[1]) ** 2)

    def execute_instructions_sequence(self) -> bool:
        """执行整段工艺任务序列"""
        instructions = self.config.get_instructions()
        self._log("Sequence", f"开始执行全流程作业，共 {len(instructions)} 道工序")

        for i, inst in enumerate(instructions):
            if not self.execute_instruction(inst):
                self._log("Sequence", f"工序执行异常终止", "ERROR")
                return False
            if i < len(instructions) - 1:
                time.sleep(0.5)

        self._log("Sequence", "单循环全工序执行完成")
        return True

    def get_stats(self) -> Dict[str, float]:
        return self.stats.copy()

    def _log(self, context: str, message: str, level: str = "INFO"):
        if level in ["INFO", "WARN", "ERROR"]:
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            print(f"[{timestamp}] [{level}] [{context}] {message}")


# ============================================================================
# 任务总控程序
# ============================================================================

class RobotTaskController:
    """仓储AGV取放货循环任务总控"""

    def __init__(self, robot: KuboOrder, robot_id: str,
                 task_config: TaskConfigLoader, word_path: str):
        self.robot = robot
        self.robot_id = robot_id
        self.config = task_config
        self.word_path = word_path

        self.control = task_config.get_control_logic()
        self.executor = RobotCommandExecutor(robot, robot_id, task_config)
        self.doc = self._init_word()

        self.start_time = datetime.now()
        self.last_word_time = self.start_time
        self.last_battery_check = self.start_time

        self._running = True
        self._emergency_stop = False

    def _init_word(self) -> Document:
        """初始化运维报表文档"""
        doc_dir = os.path.dirname(self.word_path)
        if doc_dir and not os.path.exists(doc_dir):
            os.makedirs(doc_dir)

        if os.path.exists(self.word_path):
            return Document(self.word_path)

        doc = Document()
        title = doc.add_heading('AGV货架取放货作业运行报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        metadata = self.config.get_metadata()
        doc.add_paragraph(f'任务名称: {metadata["name"]}')
        doc.add_paragraph(f'任务描述: {metadata["description"]}')
        doc.add_paragraph(f'机器人ID: {self.robot_id}')
        doc.add_paragraph(f'程序启动时间: {self._timestamp()}')
        doc.add_paragraph('=' * 70)
        doc.add_heading('周期运行日志', level=1)

        doc.save(self.word_path)
        return doc

    def _timestamp(self) -> str:
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    def check_runtime(self) -> bool:
        """运行时长安全校验"""
        elapsed = (datetime.now() - self.start_time).total_seconds() / 3600
        if elapsed >= self.control['max_runtime_hours']:
            self._log("Runtime", f"到达设定运行时长上限")
            return False
        return True

    def check_battery(self) -> bool:
        """电池电量巡检与低电量回充策略"""
        now = datetime.now()
        check_interval = self.control.get('check_interval', 10)

        if (now - self.last_battery_check).total_seconds() < check_interval:
            return True

        self.last_battery_check = now
        try:
            info = self.robot.get_info(self.robot_id)
            power = info.get('batteryInfo', {}).get('powerLevel', 50)

            if power <= self.control['battery_emergency_threshold']:
                self._log("Battery", f"电量极低紧急停机", "CRITICAL")
                self._emergency_stop = True
                return False

            if power < self.control['battery_low_threshold']:
                self._log("Battery", f"低电量触发自动回充", "WARN")
                if not self._do_charge():
                    return False

            return True
        except Exception as e:
            self._log("Battery", f"电量巡检异常", "ERROR")
            return True

    def _do_charge(self) -> bool:
        """自动前往充电站补能流程"""
        charge_point = self.config.get_location_coord('charging_point')
        self._log("Charge", "执行前往充电站动作")
        if not self.executor._do_move(charge_point, 600.0, 3, 1.0):
            return False

        target = self.control['battery_high_threshold']
        timeout = 240 * 60
        start = time.time()
        self._log("Charge", f"开始恒流补能，目标电量{target}%")

        while True:
            time.sleep(30)
            try:
                info = self.robot.get_info(self.robot_id)
                power = info.get('batteryInfo', {}).get('powerLevel', 0)
                if power >= target:
                    self._log("Charge", f"补能完成")
                    self.executor.stats['charge_count'] += 1
                    break
                if time.time() - start > timeout:
                    self._log("Charge", "充电流程超时", "ERROR")
                    return False
            except:
                pass

        work_point = self.config.get_location_coord('working_point')
        self._log("Charge", "返回作业待命位")
        return self.executor._do_move(work_point, 600.0, 0, 1.0)

    def check_error(self) -> bool:
        """机器人设备故障巡检"""
        try:
            info = self.robot.get_info(self.robot_id)
            errors = info.get('errorState', [])
            if errors:
                self._log("Error", f"设备故障告警: {errors[0]}", "ERROR")
                return False
            return True
        except:
            return True

    def update_word(self):
        """周期更新运维产能报表"""
        now = datetime.now()
        interval = timedelta(hours=self.control['word_log_interval_hours'])
        if now - self.last_word_time < interval:
            return

        stats = self.executor.get_stats()
        elapsed = (now - self.start_time).total_seconds() / 3600
        self.doc.add_heading(f'周期记录时间: {self._timestamp()}', level=2)

        table = self.doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        data = [
            ("累计运行时长", f"{elapsed:.2f}小时"),
            ("累计入库上架", f"{stats['pick_count']}托"),
            ("累计出库下架", f"{stats['place_count']}托"),
            ("底盘行驶里程", f"{stats['chassis_distance']:.2f}米"),
            ("自动补能次数", f"{stats['charge_count']}次"),
            ("完成作业工序", f"{stats['cycles']}道"),
        ]

        for i, (k, v) in enumerate(data):
            table.rows[i].cells[0].text = k
            table.rows[i].cells[1].text = v

        self.doc.add_paragraph('-' * 70)
        self.doc.save(self.word_path)
        self.last_word_time = now
        self._log("Word", "运维报表已同步更新")

    def run(self):
        """主循环调度入口"""
        self._log("Main", "=" * 70)
        self._log("Main", "AGV货架循环搬运任务启动")
        self._log("Main", f"设定总运行时长: {self.control['max_runtime_hours']}小时")
        self._log("Main", "=" * 70)
        metadata = self.config.get_metadata()

        try:
            while self._running and not self._emergency_stop:
                if not self.check_runtime(): break
                if not self.check_error(): break
                if not self.check_battery():
                    if self._emergency_stop: break
                    continue
                self.update_word()

                if metadata.get('loop', True):
                    if not self.executor.execute_instructions_sequence():
                        self._log("Main", "单循环作业异常", "ERROR")
                        time.sleep(5)
                        continue
                    self._log("Main", "一轮循环作业完成，进入下一轮")
                    time.sleep(self.control['cycle_pause_seconds'])
                else:
                    self.executor.execute_instructions_sequence()
                    break
        except KeyboardInterrupt:
            self._log("Main", "人工程序终止", "WARN")
        except Exception as e:
            self._log("Main", f"主控程序异常: {e}", "ERROR")

        self._shutdown()

    def _shutdown(self):
        """停机归位收尾流程"""
        self._log("Main", "执行停机归位流程")
        try:
            power = self.robot.get_info(self.robot_id).get('batteryInfo', {}).get('powerLevel', 50)
            if power < 30:
                charge_point = self.config.get_location_coord('charging_point')
                self.executor._do_move(charge_point, 600.0, 3, 1.0)
            else:
                work_point = self.config.get_location_coord('working_point')
                self.executor._do_move(work_point, 600.0, 0, 1.0)
        except:
            pass
        self._final_report()
        self._log("Main", "程序全流程结束")

    def _final_report(self):
        """生成最终项目运行总结报告"""
        stats = self.executor.get_stats()
        total_hours = (datetime.now() - self.start_time).total_seconds() / 3600

        self.doc.add_page_break()
        self.doc.add_heading("项目运行最终总结", level=1)
        table = self.doc.add_table(rows=7, cols=2)
        table.style = 'Medium Grid 1 Accent 1'

        data = [
            ("总运行时长", f"{total_hours:.2f}小时"),
            ("总上架取货量", f"{stats['pick_count']}托"),
            ("总下架放货量", f"{stats['place_count']}托"),
            ("全流程行驶里程", f"{stats['chassis_distance']:.2f}米"),
            ("自动补能总次数", f"{stats['charge_count']}次"),
            ("累计执行作业工序", f"{stats['cycles']}道"),
            ("最终运行状态", "正常完成" if not self._emergency_stop else "异常停机"),
        ]

        for i, (k, v) in enumerate(data):
            table.rows[i].cells[0].text = k
            table.rows[i].cells[1].text = v
            if i == len(data) - 1:
                for cell in table.rows[i].cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.bold = True
                            r.font.color.rgb = RGBColor(0, 128, 0) if not self._emergency_stop else RGBColor(255, 0, 0)

        self.doc.add_paragraph(f'报告生成时间: {self._timestamp()}')
        self.doc.save(self.word_path)

        print("\n" + "=" * 70)
        print("项目运行总览:")
        print(f"  取货上架: {stats['pick_count']}托")
        print(f"  放货下架: {stats['place_count']}托")
        print(f"  行驶总里程: {stats['chassis_distance']:.2f}米")
        print(f"  运维报告路径: {self.word_path}")
        print("=" * 70)

    def _log(self, context: str, message: str, level: str = "INFO"):
        if level in ["INFO", "WARN", "ERROR", "CRITICAL"]:
            print(f"[{self._timestamp()}] [{level}] [{context}] {message}")


# ============================================================================
# 程序启动入口
# ============================================================================
def main():
    import argparse
    parser = argparse.ArgumentParser(description='AGV仓储取放货任务控制器')
    parser.add_argument('--config', default=SystemConfig.DEFAULT_JSON_PATH)
    parser.add_argument('--word', default=SystemConfig.DEFAULT_WORD_PATH)
    args = parser.parse_args()

    print("=" * 70)
    print("AGV货架循环搬运任务控制系统")
    print(f"工艺配置路径: {args.config}")
    print("=" * 70)

    try:
        config = TaskConfigLoader(args.config)
        meta = config.get_metadata()
        print(f"任务方案: {meta['name']}")
        print(f"方案描述: {meta['description']}")
        print(f"总作业工序: {len(config.get_instructions())}道")
    except Exception as e:
        print(f"[ERROR] 工艺配置解析失败: {e}")
        sys.exit(1)

    robot = KuboOrder(SystemConfig.IP, SystemConfig.PORT, timeout=120)
    try:
        robot.robot_init(SystemConfig.ROBOT_ID)
        time.sleep(2)
        print("[INFO] 机器人底盘初始化完成")
    except Exception as e:
        print(f"[WARN] 底盘通讯初始化提示: {e}")

    controller = RobotTaskController(robot, SystemConfig.ROBOT_ID, config, args.word)
    controller.run()
    sys.exit(0)


if __name__ == "__main__":
    main()